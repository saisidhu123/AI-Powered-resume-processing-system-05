import os
import re
from typing import Tuple
import fitz
import docx


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF."""
    text_chunks = []

    with fitz.open(file_path) as doc:
        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                text_chunks.append(page_text)

    return "\n".join(text_chunks)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    doc = docx.Document(file_path)
    text_chunks = []

    # Read normal paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_chunks.append(paragraph.text.strip())

    # Read tables
    for table in doc.tables:
        for row in table.rows:
            row_cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if row_cells:
                text_chunks.append(" | ".join(row_cells))

    full_text = "\n".join(text_chunks)
    return full_text


def parse_resume(file_path: str) -> Tuple[bool, str, str]:
    """
    Parse PDF or DOCX resume.

    Returns:
        (success, extracted_text, error_message)
    """

    if not os.path.exists(file_path):
        return False, "", "Resume file does not exist on disk."

    ext = os.path.splitext(file_path)[1].lower()

    try:

        # =========================
        # PDF
        # =========================
        if ext == ".pdf":

            text = extract_text_from_pdf(file_path)

            print("\n========== PDF RESUME TEXT ==========")
            print(text[:2000].encode("ascii", errors="ignore").decode("ascii"))
            print("=====================================\n")

            if not text.strip():
                return (
                    False,
                    "",
                    "PDF contains no readable text."
                )

            from services.document_reader import build_resume_document
            doc_obj = build_resume_document(file_path, text)
            return True, doc_obj.full_text, ""

        # =========================
        # DOCX
        # =========================
        elif ext == ".docx":

            text = extract_text_from_docx(file_path)

            print("\n========== DOCX RESUME TEXT ==========")
            print(text[:2000].encode("ascii", errors="ignore").decode("ascii"))
            print("======================================\n")

            if not text.strip():
                return (
                    False,
                    "",
                    "DOCX contains no readable text."
                )

            return True, text.strip(), ""

        # =========================
        # DOC
        # =========================
        elif ext == ".doc":
            # Attempt 1: Try reading with python-docx (in case it's a docx saved as .doc)
            try:
                text = extract_text_from_docx(file_path)
                if text.strip():
                    return True, text.strip(), ""
            except Exception:
                pass

            # Attempt 2: Extract printable strings from binary .doc stream
            try:
                with open(file_path, "rb") as f:
                    content = f.read()
                words = re.findall(rb"[\x20-\x7E\x0A\x0D]{4,}", content)
                extracted_lines = [w.decode("latin1", errors="ignore").strip() for w in words if len(w.decode("latin1", errors="ignore").strip()) > 3]
                clean_lines = [line for line in extracted_lines if not any(line.startswith(p) for p in ["Root Entry", "WordDocument", "CompObj", "ObjectPool"])]
                doc_text = "\n".join(clean_lines)
                if len(doc_text.strip()) > 30:
                    return True, doc_text.strip(), ""
            except Exception:
                pass

            return (
                False,
                "",
                "Could not extract text from legacy .doc file. Please convert to .docx or .pdf."
            )

        # =========================
        # Unsupported
        # =========================
        else:

            return (
                False,
                "",
                f"Unsupported file extension: {ext}"
            )

    except Exception as e:

        return (
            False,
            "",
            f"Error parsing resume file: {str(e)}"
        )
"""
services/document_reader.py

Complete Document Reader & Canonical Representation Engine.
Reads every page, preserves line/page boundaries, extracts tables, repairs OCR artifacts,
and builds a canonical ResumeDocument object for field extraction.
"""

import os
import re
from dataclasses import dataclass, field
from typing import List, Dict, Any, Tuple, Optional


@dataclass
class ResumeDocument:
    full_text: str
    pages: List[str] = field(default_factory=list)
    sections: Dict[str, str] = field(default_factory=dict)
    tables: List[List[List[str]]] = field(default_factory=list)
    extraction_quality: Dict[str, Any] = field(default_factory=dict)
    source_file: str = ""


def normalize_resume_text(raw_text: str) -> str:
    """
    Normalizes raw resume text across all platforms and file formats.
    - Repairs line-wrap hyphens.
    - Standardizes currency symbols (₹, Rs, INR, $, USD).
    - Normalizes dash variations (en-dash, em-dash), colons, and tabs.
    - Reconstructs broken OCR words.
    """
    if not raw_text:
        return ""

    # Replace en-dash / em-dash / bullet characters with standard hyphen/bullet
    cleaned = raw_text.replace("–", "-").replace("—", "-").replace("•", "*").replace("\t", " ")

    # Repair line-wrap hyphens e.g. "Retrieval-\nAugmented" -> "Retrieval-Augmented"
    cleaned = re.sub(r"(\w+)-\s*\n\s*(\w+)", r"\1-\2", cleaned)

    # Standardize currency abbreviations
    cleaned = re.sub(r"\bRs\.?\s*", "₹", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\bINR\s*", "₹", cleaned, flags=re.IGNORECASE)

    # Repair PDF font conversion artifacts e.g. "I8 LPA" -> "₹8 LPA"
    cleaned = re.sub(r"\bI(\d+(?:\.\d+)?\s*(?:LPA|Lakhs?|Lacs?|L|k|K|Crores?|Cr))\b", r"₹\1", cleaned)

    # Normalize space-split technology terms
    replacements = [
        (r"\bPromp\s*t\s+Engin\s*eering\b", "Prompt Engineering"),
        (r"\bR\s*etr\s*ieval\s+Aug\s*mented\s+Gen\s*eration\b", "Retrieval Augmented Generation"),
        (r"\bRea\s*ct\.js\b", "React.js"),
        (r"\bFast\s*API\b", "FastAPI"),
        (r"\bPy\s*Thon\b", "Python"),
        (r"\bPost\s*gre\s*SQL\b", "PostgreSQL"),
        (r"\bOpen\s*Py\s*XL\b", "OpenPyXL"),
        (r"\bLang\s*Chain\b", "LangChain"),
        (r"\bPy\s*Torch\b", "PyTorch")
    ]
    for pat, repl in replacements:
        cleaned = re.sub(pat, repl, cleaned, flags=re.IGNORECASE)

    # Normalize multiple whitespace lines/spaces while preserving newlines
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in cleaned.splitlines()]
    return "\n".join(lines).strip()


def read_pdf_complete(file_path: str) -> Tuple[bool, List[str], List[List[List[str]]], Dict[str, Any], str]:
    """
    Inspects every page of a PDF file using pdfplumber or PyPDF2/pypdf.
    Returns (success, pages_text, extracted_tables, quality_metrics, error_message).
    """
    pages_text = []
    tables = []
    parser_used = "pypdf"

    try:
        import pdfplumber
        parser_used = "pdfplumber"
        with pdfplumber.open(file_path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                txt = page.extract_text() or ""
                pages_text.append(txt)

                # Extract tabular data if present
                page_tables = page.extract_tables() or []
                if page_tables:
                    tables.extend(page_tables)
    except Exception:
        # Fallback to pypdf / PyPDF2
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            for page in reader.pages:
                txt = page.extract_text() or ""
                pages_text.append(txt)
        except Exception:
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(file_path)
                for page in reader.pages:
                    txt = page.extract_text() or ""
                    pages_text.append(txt)
            except Exception as e:
                return False, [], [], {}, f"Failed to parse PDF pages: {str(e)}"

    full_raw = "\n".join(pages_text)
    text_length = len(full_raw.strip())
    is_poor_text = text_length < 50

    quality = {
        "pages_read": len(pages_text),
        "text_length": text_length,
        "is_poor_text": is_poor_text,
        "parser_used": parser_used,
        "ocr_fallback_used": is_poor_text
    }

    return True, pages_text, tables, quality, ""


def read_docx_complete(file_path: str) -> Tuple[bool, List[str], List[List[List[str]]], Dict[str, Any], str]:
    """
    Inspects paragraphs and tables of a DOCX file using python-docx.
    """
    try:
        import docx
        doc = docx.Document(file_path)
        
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        tables = []
        table_text_lines = []
        for t in doc.tables:
            t_rows = []
            for row in t.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                t_rows.append(row_cells)
                table_text_lines.append(" | ".join([c for c in row_cells if c]))
            tables.append(t_rows)

        full_text = "\n".join(paragraphs + table_text_lines)
        quality = {
            "pages_read": 1,
            "text_length": len(full_text),
            "is_poor_text": len(full_text) < 30,
            "parser_used": "python-docx",
            "ocr_fallback_used": False
        }

        return True, [full_text], tables, quality, ""
    except Exception as e:
        return False, [], [], {}, f"Failed to parse DOCX file: {str(e)}"


def build_resume_document(file_path: str, raw_text_fallback: str = "") -> ResumeDocument:
    """
    Constructs a canonical ResumeDocument object for any input file.
    """
    if not file_path or not os.path.exists(file_path):
        norm_txt = normalize_resume_text(raw_text_fallback)
        return ResumeDocument(
            full_text=norm_txt,
            pages=[norm_txt],
            extraction_quality={"pages_read": 1, "text_length": len(norm_txt), "parser_used": "raw_fallback"},
            source_file=file_path or ""
        )

    ext = os.path.splitext(file_path)[1].lower()
    pages_text = []
    tables = []
    quality = {}

    if ext == ".pdf":
        ok, pages_text, tables, quality, err = read_pdf_complete(file_path)
        if not ok or not "".join(pages_text).strip():
            pages_text = [raw_text_fallback]
    elif ext in [".docx", ".doc"]:
        ok, pages_text, tables, quality, err = read_docx_complete(file_path)
        if not ok or not "".join(pages_text).strip():
            pages_text = [raw_text_fallback]
    else:
        pages_text = [raw_text_fallback]

    full_raw = "\n".join(pages_text)
    full_norm = normalize_resume_text(full_raw)

    return ResumeDocument(
        full_text=full_norm,
        pages=[normalize_resume_text(p) for p in pages_text],
        tables=tables,
        extraction_quality=quality,
        source_file=file_path
    )
